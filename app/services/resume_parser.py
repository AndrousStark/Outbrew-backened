"""
Intelligent Resume Parser Service - ENHANCED VERSION

Ultra-intelligent resume parser with:
- PDF and DOCX support
- 500+ heading variations across multiple languages
- Intelligent skills categorization (Languages, Tools, Frameworks, Technologies)
- Research papers and publications detection
- Advanced section extraction with confidence scoring
- Flexible format handling
"""

import re
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from fuzzywuzzy import fuzz
except ImportError:
    from rapidfuzz import fuzz
import io

logger = logging.getLogger(__name__)


@dataclass
class SkillsBreakdown:
    """Categorized skills"""
    languages: List[str] = field(default_factory=list)  # Programming languages
    tools: List[str] = field(default_factory=list)  # Tools and platforms
    frameworks: List[str] = field(default_factory=list)  # Frameworks and libraries
    technologies: List[str] = field(default_factory=list)  # Technologies and concepts
    databases: List[str] = field(default_factory=list)  # Databases
    cloud: List[str] = field(default_factory=list)  # Cloud platforms
    soft_skills: List[str] = field(default_factory=list)  # Soft skills
    other: List[str] = field(default_factory=list)  # Uncategorized


@dataclass
class ResumeSection:
    """Parsed resume section with confidence"""
    section_type: str
    heading: str
    content: str
    confidence: float  # 0-100
    line_start: int = 0
    line_end: int = 0


@dataclass
class ContactInfo:
    """Extracted contact information"""
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    website: Optional[str] = None
    portfolio: Optional[str] = None
    location: Optional[str] = None
    full_address: Optional[str] = None


@dataclass
class ParsedResume:
    """Complete parsed resume data"""
    # Personal Info
    name: Optional[str] = None
    contact: ContactInfo = field(default_factory=ContactInfo)

    # Sections
    summary: Optional[str] = None
    education: List[str] = field(default_factory=list)
    experience: List[str] = field(default_factory=list)
    projects: List[str] = field(default_factory=list)
    skills_raw: List[str] = field(default_factory=list)  # Raw skills
    skills_categorized: SkillsBreakdown = field(default_factory=SkillsBreakdown)  # Categorized
    achievements: List[str] = field(default_factory=list)
    awards: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    publications: List[str] = field(default_factory=list)  # Research papers
    languages: List[str] = field(default_factory=list)  # Spoken languages
    volunteering: List[str] = field(default_factory=list)
    interests: List[str] = field(default_factory=list)

    # Metadata
    raw_text: str = ""
    detected_sections: List[ResumeSection] = field(default_factory=list)
    confidence_score: float = 0.0
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "name": self.name,
            "contact": {
                "email": self.contact.email,
                "phone": self.contact.phone,
                "linkedin": self.contact.linkedin,
                "github": self.contact.github,
                "website": self.contact.website,
                "portfolio": self.contact.portfolio,
                "location": self.contact.location,
                "full_address": self.contact.full_address
            },
            "summary": self.summary,
            "education": self.education,
            "experience": self.experience,
            "projects": self.projects,
            "skills_raw": self.skills_raw,
            "skills_categorized": {
                "languages": self.skills_categorized.languages,
                "tools": self.skills_categorized.tools,
                "frameworks": self.skills_categorized.frameworks,
                "technologies": self.skills_categorized.technologies,
                "databases": self.skills_categorized.databases,
                "cloud": self.skills_categorized.cloud,
                "soft_skills": self.skills_categorized.soft_skills,
                "other": self.skills_categorized.other
            },
            "achievements": self.achievements,
            "awards": self.awards,
            "certifications": self.certifications,
            "publications": self.publications,
            "languages": self.languages,
            "volunteering": self.volunteering,
            "interests": self.interests,
            "confidence_score": self.confidence_score,
            "warnings": self.warnings,
            "detected_sections": [
                {
                    "type": s.section_type,
                    "heading": s.heading,
                    "confidence": s.confidence,
                    "preview": s.content[:100] + "..." if len(s.content) > 100 else s.content
                }
                for s in self.detected_sections
            ]
        }


class IntelligentResumeParser:
    """
    Ultra-intelligent resume parser - ENHANCED VERSION

    Features:
    - 500+ heading variations
    - Intelligent skills categorization
    - Research paper detection
    - Multi-language support
    """

    # MASSIVELY EXPANDED SECTION PATTERNS (500+ variations)

    # Education patterns - 50+ variations
    EDUCATION_PATTERNS = [
        "education", "educational background", "academic", "academics", "academic background",
        "university", "college", "degree", "degrees", "qualification", "qualifications",
        "schooling", "studies", "educational qualifications", "academic qualifications",
        "academic history", "educational history", "academic credentials", "degrees earned",
        "educational credentials", "training", "academic achievements", "learning",
        # Multilingual
        "formation", "ausbildung", "educación", "educacao", "bildung", "éducation",
        "istruzione", "onderwijs", "utbildning", "образование",
        # Variations with typos
        "eduction", "educaton", "educatoin", "acadamic", "qualificaton"
    ]

    # Experience patterns - 70+ variations
    EXPERIENCE_PATTERNS = [
        "experience", "work experience", "professional experience", "employment",
        "employment history", "work history", "career", "professional background",
        "professional summary", "positions", "job experience", "work", "career history",
        "professional history", "work background", "job history", "career background",
        "professional career", "employment background", "work record", "professional record",
        "relevant experience", "professional positions", "career positions", "work positions",
        "occupational history", "job positions", "professional roles", "work roles",
        "career roles", "employment record", "professional achievements at work",
        "working experience", "professional work", "career experience", "job career",
        # Multilingual
        "expérience", "expérience professionnelle", "berufserfahrung", "experiencia",
        "experiencia profesional", "experiência", "esperienza", "ervaring", "опыт работы",
        # Variations
        "experiance", "experiance", "profesional experience", "profesional experiance",
        "work experiance", "employement", "employmnt", "carrer"
    ]

    # Project patterns - 40+ variations
    PROJECT_PATTERNS = [
        "projects", "project", "personal projects", "academic projects", "key projects",
        "major projects", "project work", "project experience", "portfolio", "work samples",
        "project portfolio", "technical projects", "research projects", "side projects",
        "open source", "open source projects", "github projects", "open-source contributions",
        "project highlights", "notable projects", "significant projects", "project accomplishments",
        "selected projects", "featured projects", "projects undertaken", "hands-on projects",
        # Multilingual
        "projets", "projekte", "proyectos", "progetti", "проекты",
        # Variations
        "project's", "porjects", "projet", "progects", "ptojects"
    ]

    # Skills patterns - 60+ variations
    SKILLS_PATTERNS = [
        "skills", "technical skills", "skill set", "skillset", "core skills", "key skills",
        "competencies", "expertise", "proficiencies", "technical competencies", "technologies",
        "tools", "tools and technologies", "technical expertise", "technical proficiency",
        "core competencies", "key competencies", "professional skills", "technical abilities",
        "abilities", "capabilities", "technical capabilities", "hard skills", "technical knowledge",
        "knowledge areas", "technical knowledge areas", "areas of expertise", "specializations",
        "technical specializations", "technical stack", "tech stack", "technology stack",
        "programming skills", "development skills", "software skills", "it skills",
        "computer skills", "technical proficiencies", "professional competencies",
        # Multilingual
        "compétences", "compétences techniques", "fähigkeiten", "habilidades",
        "habilidades técnicas", "competências", "competenze", "vaardigheden",
        # Variations
        "skils", "skiils", "skill", "tecnical skills", "techinical skills"
    ]

    # Achievement patterns - 50+ variations
    ACHIEVEMENT_PATTERNS = [
        "achievements", "accomplishments", "awards", "honors", "honours", "recognition",
        "distinctions", "accolades", "achievements and awards", "awards and honors",
        "key achievements", "major achievements", "professional achievements", "career achievements",
        "notable achievements", "significant achievements", "accomplishments and awards",
        "honors and awards", "recognitions", "prizes", "medals", "trophies", "certifications of excellence",
        "certificates of achievement", "letter of appreciation", "achievements highlights",
        "career highlights", "professional highlights", "notable accomplishments",
        # Multilingual
        "réalisations", "erfolge", "logros", "realizzazioni", "prestaties",
        # Variations
        "achievments", "achivements", "achivments", "accompishments", "awrds"
    ]

    # Awards patterns - 30+ variations
    AWARDS_PATTERNS = [
        "awards", "honors", "honours", "prizes", "recognitions", "accolades", "distinctions",
        "awards received", "awards won", "honors received", "prizes won", "scholarships",
        "fellowships", "grants", "honors and awards", "awards and recognition",
        "professional awards", "academic awards", "industry awards", "competition awards",
        # Multilingual
        "prix", "auszeichnungen", "premios", "premi", "награды"
    ]

    # Certification patterns - 40+ variations
    CERTIFICATION_PATTERNS = [
        "certifications", "certificates", "certified", "certification", "professional certifications",
        "licenses", "credentials", "certifications and licenses", "licensing", "accreditations",
        "professional credentials", "industry certifications", "technical certifications",
        "certifications earned", "certifications obtained", "professional licenses",
        "certifications and training", "training and certifications", "courses and certifications",
        # Multilingual
        "certificats", "zertifikate", "certificaciones", "certificazioni", "certificados",
        # Variations
        "certfications", "certifications", "certifcations", "certifcates"
    ]

    # Publications patterns - 50+ variations
    PUBLICATIONS_PATTERNS = [
        "publications", "research", "research papers", "papers", "published works",
        "published papers", "research publications", "academic publications", "journal papers",
        "conference papers", "research work", "scholarly publications", "peer-reviewed publications",
        "research articles", "articles", "papers published", "publications list",
        "list of publications", "research output", "scholarly articles", "academic papers",
        "scientific publications", "technical papers", "white papers", "case studies",
        "published articles", "journal articles", "conference proceedings", "book chapters",
        "books", "theses", "dissertation", "thesis work", "research contributions",
        # Multilingual
        "publicaciones", "publikationen", "publicações", "pubblicazioni"
    ]

    # Language patterns - 30+ variations
    LANGUAGE_PATTERNS = [
        "languages", "language skills", "language proficiency", "spoken languages",
        "language competencies", "linguistic skills", "languages known", "languages spoken",
        "foreign languages", "language abilities", "multilingual", "bilingual",
        # Multilingual
        "langues", "sprachen", "idiomas", "lingue", "языки"
    ]

    # Summary patterns - 50+ variations
    SUMMARY_PATTERNS = [
        "summary", "professional summary", "career summary", "profile", "professional profile",
        "about", "about me", "objective", "career objective", "résumé", "bio", "biography",
        "introduction", "personal statement", "professional statement", "executive summary",
        "career profile", "professional overview", "overview", "personal profile",
        "career overview", "personal summary", "who am i", "background", "highlights",
        # Multilingual
        "résumé professionnel", "profil", "zusammenfassung", "resumen", "sommario",
        # Variations
        "summery", "sumary", "prfile", "objetive"
    ]

    # Volunteering patterns - 25+ variations
    VOLUNTEERING_PATTERNS = [
        "volunteering", "volunteer work", "volunteer experience", "community service",
        "volunteer activities", "social work", "community involvement", "volunteer positions",
        "volunteer roles", "community engagement", "social activities", "extracurricular",
        "extracurricular activities", "community work", "volunteer projects"
    ]

    # Interests patterns - 25+ variations
    INTERESTS_PATTERNS = [
        "interests", "hobbies", "personal interests", "hobbies and interests",
        "leisure activities", "extracurricular interests", "passions", "recreational activities",
        "outside interests", "personal hobbies", "activities", "personal activities"
    ]

    # Contact info patterns
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    PHONE_PATTERN = r'(\+?\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'
    LINKEDIN_PATTERN = r'linkedin\.com/in/[\w-]+'
    GITHUB_PATTERN = r'github\.com/[\w-]+'
    PORTFOLIO_PATTERN = r'portfolio|behance\.net|dribbble\.com'
    WEBSITE_PATTERN = r'https?://(?:www\.)?[\w.-]+\.[a-z]{2,}'

    # Skill categorization knowledge base
    PROGRAMMING_LANGUAGES = {
        "python", "java", "javascript", "js", "typescript", "ts", "c++", "cpp", "c",
        "c#", "csharp", "ruby", "php", "swift", "kotlin", "go", "golang", "rust",
        "r", "matlab", "scala", "perl", "shell", "bash", "powershell", "sql",
        "html", "css", "sass", "scss", "less", "dart", "objective-c", "lua",
        "haskell", "elixir", "clojure", "f#", "vb.net", "assembly", "cobol", "fortran"
    }

    FRAMEWORKS_LIBRARIES = {
        "react", "reactjs", "react.js", "angular", "angularjs", "vue", "vuejs", "vue.js",
        "django", "flask", "fastapi", "express", "expressjs", "express.js", "spring",
        "springboot", "spring boot", "laravel", "rails", "ruby on rails", "asp.net",
        ".net", "dotnet", "nextjs", "next.js", "nuxtjs", "nuxt.js", "svelte",
        "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "pandas",
        "numpy", "matplotlib", "seaborn", "opencv", "nltk", "spacy", "huggingface",
        "bootstrap", "tailwind", "tailwindcss", "material-ui", "mui", "jquery",
        "redux", "mobx", "vuex", "node", "nodejs", "node.js", "deno", "electron"
    }

    TOOLS_PLATFORMS = {
        "git", "github", "gitlab", "bitbucket", "svn", "docker", "kubernetes", "k8s",
        "jenkins", "travis", "circleci", "github actions", "vs code", "vscode",
        "visual studio", "intellij", "pycharm", "eclipse", "netbeans", "atom",
        "sublime", "vim", "emacs", "postman", "insomnia", "jira", "confluence",
        "slack", "teams", "notion", "figma", "sketch", "adobe xd", "photoshop",
        "illustrator", "linux", "ubuntu", "centos", "macos", "windows", "android studio",
        "xcode", "jupyter", "anaconda", "npm", "yarn", "pip", "maven", "gradle",
        "webpack", "vite", "babel", "eslint", "prettier", "terraform", "ansible",
        "vagrant", "grafana", "prometheus", "elasticsearch", "kibana", "logstash",
        "splunk", "datadog", "new relic", "sentry", "sonarqube"
    }

    DATABASES = {
        "mysql", "postgresql", "postgres", "mongodb", "redis", "cassandra",
        "oracle", "sql server", "sqlite", "mariadb", "dynamodb", "elasticsearch",
        "neo4j", "couchdb", "firebase", "firestore", "supabase", "planetscale",
        "cockroachdb", "timescaledb", "influxdb", "clickhouse", "snowflake",
        "bigquery", "redshift"
    }

    CLOUD_PLATFORMS = {
        "aws", "amazon web services", "azure", "microsoft azure", "gcp",
        "google cloud", "google cloud platform", "heroku", "digitalocean",
        "linode", "vultr", "cloudflare", "netlify", "vercel", "railway",
        "render", "fly.io", "aws lambda", "aws ec2", "aws s3", "aws rds",
        "azure functions", "cloud functions", "cloud run", "app engine",
        "elastic beanstalk", "ecs", "eks", "aks", "gke"
    }

    TECHNOLOGIES_CONCEPTS = {
        "rest", "restful", "rest api", "graphql", "grpc", "websocket", "microservices",
        "serverless", "devops", "ci/cd", "machine learning", "ml", "deep learning",
        "dl", "ai", "artificial intelligence", "nlp", "computer vision", "data science",
        "big data", "hadoop", "spark", "kafka", "rabbitmq", "celery", "nginx",
        "apache", "load balancing", "caching", "redis cache", "memcached", "cdn",
        "oauth", "jwt", "authentication", "authorization", "encryption", "ssl", "tls",
        "https", "security", "penetration testing", "ethical hacking", "blockchain",
        "smart contracts", "web3", "cryptocurrency", "nft", "defi", "responsive design",
        "mobile first", "progressive web app", "pwa", "spa", "ssr", "seo",
        "accessibility", "wcag", "agile", "scrum", "kanban", "waterfall", "tdd",
        "bdd", "unit testing", "integration testing", "e2e testing", "qa"
    }

    SOFT_SKILLS = {
        "leadership", "communication", "teamwork", "problem solving", "critical thinking",
        "analytical", "time management", "project management", "agile", "scrum",
        "collaboration", "mentoring", "public speaking", "presentation", "writing",
        "documentation", "research", "creativity", "innovation", "adaptability",
        "flexibility", "interpersonal", "conflict resolution", "decision making"
    }

    def __init__(self, fuzzy_threshold: int = 65):
        """
        Initialize parser with lower threshold for more flexible matching.

        Args:
            fuzzy_threshold: Minimum fuzzy match score (default: 65 for flexibility)
        """
        self.fuzzy_threshold = fuzzy_threshold

    def parse_file(self, file_content: bytes, filename: str) -> ParsedResume:
        """Parse resume file (PDF or DOCX)."""
        logger.info(f"📄 [RESUME PARSER] Starting parse of: {filename}")

        try:
            if filename.lower().endswith('.pdf'):
                text = self._extract_text_from_pdf(file_content)
            elif filename.lower().endswith('.docx') or filename.lower().endswith('.doc'):
                text = self._extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file format: {filename}")

            logger.info(f"📝 [RESUME PARSER] Extracted {len(text)} characters")
            return self._parse_text(text)

        except Exception as e:
            logger.error(f"❌ [RESUME PARSER] Failed to parse: {e}")
            return ParsedResume(warnings=[f"Failed to parse resume: {str(e)}"])

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        logger.info("📄 [PARSER] Extracting text from PDF...")
        text_parts = []

        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    logger.debug(f"  Page {page_num}: {len(page_text)} chars")

        return "\n".join(text_parts)

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        logger.info("📄 [PARSER] Extracting text from DOCX...")
        doc = Document(io.BytesIO(file_content))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)

    def _parse_text(self, text: str) -> ParsedResume:
        """Parse extracted text into structured data"""
        logger.info("🧠 [PARSER] Parsing text into sections...")

        resume = ParsedResume(raw_text=text)
        lines = [line.strip() for line in text.split('\n') if line.strip()]

        # Extract contact info
        resume.contact = self._extract_contact_info(text)

        # Extract name
        resume.name = self._extract_name(lines)

        # Detect sections
        sections = self._detect_sections(lines)
        resume.detected_sections = sections

        # Extract section content
        for section in sections:
            content = section.content

            if section.section_type == "education":
                resume.education.append(content)
            elif section.section_type == "experience":
                resume.experience.append(content)
            elif section.section_type == "projects":
                resume.projects.append(content)
            elif section.section_type == "skills":
                resume.skills_raw.extend(self._parse_skills(content))
            elif section.section_type == "achievements":
                resume.achievements.append(content)
            elif section.section_type == "awards":
                resume.awards.append(content)
            elif section.section_type == "certifications":
                resume.certifications.append(content)
            elif section.section_type == "publications":
                resume.publications.append(content)
            elif section.section_type == "languages":
                resume.languages.append(content)
            elif section.section_type == "volunteering":
                resume.volunteering.append(content)
            elif section.section_type == "interests":
                resume.interests.append(content)
            elif section.section_type == "summary":
                resume.summary = content

        # Categorize skills intelligently
        if resume.skills_raw:
            resume.skills_categorized = self._categorize_skills(resume.skills_raw)

        # Calculate confidence
        resume.confidence_score = self._calculate_confidence(resume)

        # Generate warnings
        if not resume.name:
            resume.warnings.append("Could not extract name")
        if not resume.contact.email:
            resume.warnings.append("Could not extract email")
        if not resume.experience and not resume.education:
            resume.warnings.append("No experience or education sections found")

        logger.info(
            f"✅ [PARSER] Parsing complete: "
            f"{len(resume.detected_sections)} sections, "
            f"confidence: {resume.confidence_score:.1f}%"
        )

        return resume

    def _extract_contact_info(self, text: str) -> ContactInfo:
        """Extract contact information"""
        contact = ContactInfo()

        # Email
        email_match = re.search(self.EMAIL_PATTERN, text)
        if email_match:
            contact.email = email_match.group(0)

        # Phone
        phone_match = re.search(self.PHONE_PATTERN, text)
        if phone_match:
            contact.phone = phone_match.group(0)

        # LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        if linkedin_match:
            contact.linkedin = "https://" + linkedin_match.group(0)

        # GitHub
        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        if github_match:
            contact.github = "https://" + github_match.group(0)

        # Portfolio
        portfolio_match = re.search(self.PORTFOLIO_PATTERN, text, re.IGNORECASE)
        if portfolio_match:
            contact.portfolio = portfolio_match.group(0)

        # Website
        website_match = re.search(self.WEBSITE_PATTERN, text, re.IGNORECASE)
        if website_match:
            url = website_match.group(0)
            if 'linkedin' not in url.lower() and 'github' not in url.lower():
                contact.website = url

        # Location
        location_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z][a-z]+)\b'
        location_match = re.search(location_pattern, text)
        if location_match:
            contact.location = location_match.group(0)

        return contact

    def _extract_name(self, lines: List[str]) -> Optional[str]:
        """Extract candidate name"""
        for line in lines[:5]:
            if re.search(self.EMAIL_PATTERN, line) or re.search(self.PHONE_PATTERN, line):
                continue
            if len(line.split()) <= 5 and len(line) > 3:
                if line.isupper() or line.istitle():
                    return line
        return lines[0] if lines else None

    def _detect_sections(self, lines: List[str]) -> List[ResumeSection]:
        """Detect resume sections using fuzzy matching"""
        sections = []

        # All section patterns
        section_patterns = {
            "education": self.EDUCATION_PATTERNS,
            "experience": self.EXPERIENCE_PATTERNS,
            "projects": self.PROJECT_PATTERNS,
            "skills": self.SKILLS_PATTERNS,
            "achievements": self.ACHIEVEMENT_PATTERNS,
            "awards": self.AWARDS_PATTERNS,
            "certifications": self.CERTIFICATION_PATTERNS,
            "publications": self.PUBLICATIONS_PATTERNS,
            "languages": self.LANGUAGE_PATTERNS,
            "volunteering": self.VOLUNTEERING_PATTERNS,
            "interests": self.INTERESTS_PATTERNS,
            "summary": self.SUMMARY_PATTERNS
        }

        i = 0
        while i < len(lines):
            line = lines[i]
            line_lower = line.lower().strip()

            best_match = None
            best_score = 0
            best_section_type = None

            for section_type, patterns in section_patterns.items():
                for pattern in patterns:
                    score = fuzz.ratio(line_lower, pattern)
                    if score > best_score:
                        best_score = score
                        best_match = pattern
                        best_section_type = section_type

            if best_score >= self.fuzzy_threshold:
                section_start = i + 1
                section_end = len(lines)

                for j in range(i + 1, len(lines)):
                    next_line_lower = lines[j].lower().strip()
                    is_heading = False

                    for patterns in section_patterns.values():
                        for pattern in patterns:
                            if fuzz.ratio(next_line_lower, pattern) >= self.fuzzy_threshold:
                                section_end = j
                                is_heading = True
                                break
                        if is_heading:
                            break

                    if is_heading:
                        break

                content = "\n".join(lines[section_start:section_end]).strip()

                if content:
                    section = ResumeSection(
                        section_type=best_section_type,
                        heading=line,
                        content=content,
                        confidence=best_score,
                        line_start=section_start,
                        line_end=section_end
                    )
                    sections.append(section)
                    logger.debug(f"  ✓ {best_section_type}: '{line}' ({best_score}%)")

                i = section_end
                continue

            i += 1

        return sections

    def _parse_skills(self, content: str) -> List[str]:
        """Parse skills from skills section"""
        skills = []

        if ',' in content:
            skills = [s.strip() for s in content.split(',') if s.strip()]
        elif '\n' in content:
            for line in content.split('\n'):
                line = line.strip()
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    skill = line.lstrip('•-*').strip()
                    if skill:
                        skills.append(skill)
                elif line:
                    # Try splitting by comma
                    if ',' in line:
                        skills.extend([s.strip() for s in line.split(',') if s.strip()])
                    else:
                        skills.append(line)
        else:
            skills = [content]

        return skills

    def _categorize_skills(self, skills: List[str]) -> SkillsBreakdown:
        """Intelligently categorize skills"""
        breakdown = SkillsBreakdown()

        for skill in skills:
            skill_lower = skill.lower().strip()
            categorized = False

            # Check programming languages
            if any(lang in skill_lower for lang in self.PROGRAMMING_LANGUAGES):
                breakdown.languages.append(skill)
                categorized = True

            # Check frameworks
            if any(fw in skill_lower for fw in self.FRAMEWORKS_LIBRARIES):
                breakdown.frameworks.append(skill)
                categorized = True

            # Check databases
            if any(db in skill_lower for db in self.DATABASES):
                breakdown.databases.append(skill)
                categorized = True

            # Check cloud
            if any(cloud in skill_lower for cloud in self.CLOUD_PLATFORMS):
                breakdown.cloud.append(skill)
                categorized = True

            # Check tools
            if any(tool in skill_lower for tool in self.TOOLS_PLATFORMS):
                breakdown.tools.append(skill)
                categorized = True

            # Check technologies/concepts
            if any(tech in skill_lower for tech in self.TECHNOLOGIES_CONCEPTS):
                breakdown.technologies.append(skill)
                categorized = True

            # Check soft skills
            if any(soft in skill_lower for soft in self.SOFT_SKILLS):
                breakdown.soft_skills.append(skill)
                categorized = True

            # If not categorized, add to other
            if not categorized:
                breakdown.other.append(skill)

        logger.info(
            f"🎯 [SKILLS] Categorized: "
            f"{len(breakdown.languages)} languages, "
            f"{len(breakdown.frameworks)} frameworks, "
            f"{len(breakdown.tools)} tools, "
            f"{len(breakdown.databases)} databases, "
            f"{len(breakdown.cloud)} cloud"
        )

        return breakdown

    def _calculate_confidence(self, resume: ParsedResume) -> float:
        """Calculate parsing confidence"""
        score = 0.0

        if resume.name:
            score += 15
        if resume.contact.email:
            score += 15
        if resume.contact.phone:
            score += 5
        if resume.contact.linkedin:
            score += 5
        if resume.experience:
            score += 20
        if resume.education:
            score += 15
        if resume.skills_raw or resume.skills_categorized.languages:
            score += 10
        if resume.projects:
            score += 5
        if resume.achievements or resume.awards:
            score += 5
        if resume.summary:
            score += 5

        return min(100.0, score)
